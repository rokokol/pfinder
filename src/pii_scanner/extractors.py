from __future__ import annotations

import html
import logging
import re
import zipfile
from csv import DictReader, Sniffer
from dataclasses import dataclass, field
from io import BytesIO, StringIO
from pathlib import Path


@dataclass
class ExtractionResult:
    text: str
    file_format: str
    warnings: list[str]
    scan_texts: list[str] = field(default_factory=list)


TEXT_EXTENSIONS = {".txt", ".md", ".csv", ".json", ".xml", ".log"}
HTML_EXTENSIONS = {".html", ".htm"}
IMAGE_EXTENSIONS = {".tif", ".tiff", ".jpg", ".jpeg", ".png", ".gif"}
SUPPORTED_BINARY_EXTENSIONS = {".pdf", ".doc", ".docx", ".rtf", ".parquet", ".xls", ".xlsx", *IMAGE_EXTENSIONS}
OCR_LANGUAGE_ALIASES = {
    "ru": "rus",
    "rus": "rus",
    "russian": "rus",
    "en": "eng",
    "eng": "eng",
    "english": "eng",
}


def normalize_ocr_languages(languages: tuple[str, ...] | list[str] | None) -> tuple[str, ...]:
    normalized: list[str] = []
    for language in languages or ():
        value = language.strip().lower()
        if not value:
            continue
        normalized_value = OCR_LANGUAGE_ALIASES.get(value, value)
        if normalized_value not in normalized:
            normalized.append(normalized_value)
    return tuple(normalized)


def _ocr_language_argument(languages: tuple[str, ...]) -> str | None:
    if not languages:
        return None
    return "+".join(languages)


def _serial_ocr_language_specs(languages: tuple[str, ...]) -> list[tuple[str | None, str]]:
    if not languages:
        return [(None, "default")]
    specs: list[tuple[str | None, str]] = [(language, language) for language in languages]
    if len(languages) > 1:
        combined = "+".join(languages)
        specs.append((combined, combined))
    return specs


def _limit_text(text: str, max_chars: int, warnings: list[str]) -> str:
    if max_chars > 0 and len(text) > max_chars:
        warnings.append(f"text truncated to {max_chars} characters")
        return text[:max_chars]
    return text


def _read_limited_bytes(path: Path, max_bytes: int, warnings: list[str]) -> bytes:
    size = path.stat().st_size
    read_size = size if max_bytes <= 0 else min(size, max_bytes)
    if max_bytes > 0 and size > max_bytes:
        warnings.append(f"file truncated to {max_bytes} bytes before extraction")
    with path.open("rb") as handle:
        return handle.read(read_size)


def _decode_bytes(data: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "cp1251", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="ignore")


def _looks_like_text_bytes(data: bytes) -> bool:
    if not data:
        return True
    if data.startswith((b"\x7fELF", b"MZ", b"\x89PNG", b"\xff\xd8", b"GIF8", b"%PDF", b"PK\x03\x04")):
        return False
    sample = data[:8192]
    if b"\x00" in sample:
        return False
    printable = sum(1 for byte in sample if byte in b"\n\r\t" or 32 <= byte <= 126 or byte >= 128)
    return printable / max(len(sample), 1) > 0.85


def _strip_html(text: str) -> str:
    try:
        from bs4 import BeautifulSoup

        soup = BeautifulSoup(text, "html.parser")
        for tag in soup(["script", "style"]):
            tag.decompose()
        return soup.get_text("\n")
    except Exception:
        return re.sub(r"<[^>]+>", " ", html.unescape(text))


def _strip_rtf(text: str) -> str:
    try:
        from striprtf.striprtf import rtf_to_text

        return rtf_to_text(text)
    except Exception:
        return text


class _LogCapture(logging.Handler):
    def __init__(self) -> None:
        super().__init__(level=logging.WARNING)
        self.messages: list[str] = []

    def emit(self, record: logging.LogRecord) -> None:
        self.messages.append(record.getMessage())


def _combine_scan_texts(text: str, ocr_texts: list[str], serial_ocr: bool) -> tuple[str, list[str]]:
    parts = [part for part in [text, *ocr_texts] if part.strip()]
    combined_text = "\n".join(parts)
    if not serial_ocr or not ocr_texts:
        return combined_text, []
    scan_texts = [text] if text.strip() else []
    scan_texts.extend(part for part in ocr_texts if part.strip())
    return combined_text, scan_texts


def _extract_pdf(
    path: Path,
    max_pages: int,
    enable_ocr: bool,
    serial_ocr: bool,
    ocr_languages: tuple[str, ...],
    warnings: list[str],
) -> tuple[str, list[str]]:
    from pypdf import PdfReader

    pypdf_logger = logging.getLogger("pypdf")
    old_level = pypdf_logger.level
    old_propagate = pypdf_logger.propagate
    capture = _LogCapture()
    pypdf_logger.addHandler(capture)
    pypdf_logger.setLevel(logging.WARNING)
    pypdf_logger.propagate = False
    chunks: list[str] = []
    try:
        reader = PdfReader(str(path))
        if reader.is_encrypted:
            try:
                reader.decrypt("")
                warnings.append("encrypted PDF opened with empty password")
            except Exception as exc:
                warnings.append(f"encrypted PDF could not be decrypted: {exc}")
                if enable_ocr:
                    return _combine_scan_texts(
                        "",
                        _ocr_pdf(path, max_pages, warnings, serial_ocr=serial_ocr, ocr_languages=ocr_languages),
                        serial_ocr,
                    )
                return "", []
        pages = reader.pages if max_pages <= 0 else reader.pages[:max_pages]
        for index, page in enumerate(pages, start=1):
            try:
                chunks.append(page.extract_text() or "")
            except Exception as exc:
                warnings.append(f"PDF page {index} extraction failed: {exc}")
    finally:
        pypdf_logger.removeHandler(capture)
        pypdf_logger.setLevel(old_level)
        pypdf_logger.propagate = old_propagate

    if capture.messages:
        unique_messages = list(dict.fromkeys(capture.messages))
        warnings.extend(f"pypdf warning: {message}" for message in unique_messages[:5])
        if len(unique_messages) > 5:
            warnings.append(f"pypdf emitted {len(unique_messages) - 5} more warnings")

    text = "\n".join(chunks)
    needs_ocr_fallback = capture.messages or _looks_like_poor_pdf_text(text)
    if enable_ocr:
        ocr_texts: list[str] = []
        if serial_ocr:
            specs = ", ".join(label for _, label in _serial_ocr_language_specs(ocr_languages))
            warnings.append(f"serial OCR enabled for PDF: {specs}")
        if needs_ocr_fallback:
            warnings.append("PDF OCR fallback enabled after text extraction warnings or low text quality")
            ocr_texts = _ocr_pdf(path, max_pages, warnings, serial_ocr=serial_ocr, ocr_languages=ocr_languages)
        else:
            image_page_indexes = _pdf_image_page_indexes(path, max_pages, warnings)
            if image_page_indexes is None:
                warnings.append("PDF image detection unavailable; OCRing all PDF pages to cover embedded images")
                ocr_texts = _ocr_pdf(path, max_pages, warnings, serial_ocr=serial_ocr, ocr_languages=ocr_languages)
            elif image_page_indexes:
                warnings.append(f"PDF image OCR enabled for {len(image_page_indexes)} page(s) with embedded images")
                ocr_texts = _ocr_pdf(
                    path,
                    max_pages,
                    warnings,
                    page_indexes=image_page_indexes,
                    serial_ocr=serial_ocr,
                    ocr_languages=ocr_languages,
                )
        if any(part.strip() for part in ocr_texts):
            return _combine_scan_texts(text, ocr_texts, serial_ocr)
    return text, []


def _looks_like_poor_pdf_text(text: str) -> bool:
    stripped = re.sub(r"\s+", "", text)
    if not stripped:
        return True
    letters = sum(1 for char in stripped if char.isalpha())
    return len(stripped) > 80 and letters / max(len(stripped), 1) < 0.25


def _pdf_image_page_indexes(path: Path, max_pages: int, warnings: list[str]) -> list[int] | None:
    try:
        import fitz
    except Exception as exc:
        fallback = _pdf_image_page_indexes_with_pypdfium2(path, max_pages, warnings)
        if fallback is None:
            warnings.append(f"PDF image detection unavailable: {exc}")
        return fallback

    try:
        document = fitz.open(path)
    except Exception as exc:
        fallback = _pdf_image_page_indexes_with_pypdfium2(path, max_pages, warnings)
        if fallback is None:
            warnings.append(f"PDF image detection open failed: {exc}")
        return fallback

    try:
        page_count = len(document) if max_pages <= 0 else min(len(document), max_pages)
        return [index for index in range(page_count) if document.load_page(index).get_images(full=True)]
    except Exception as exc:
        fallback = _pdf_image_page_indexes_with_pypdfium2(path, max_pages, warnings)
        if fallback is None:
            warnings.append(f"PDF image detection failed: {exc}")
        return fallback
    finally:
        document.close()


def _pdf_image_page_indexes_with_pypdfium2(path: Path, max_pages: int, warnings: list[str]) -> list[int] | None:
    try:
        import pypdfium2 as pdfium
        from pypdfium2 import raw
    except Exception as exc:
        warnings.append(f"PDF image detection through pypdfium2 unavailable: {exc}")
        return None

    try:
        document = pdfium.PdfDocument(str(path))
    except Exception as exc:
        warnings.append(f"PDF image detection through pypdfium2 open failed: {exc}")
        return None

    try:
        page_count = len(document) if max_pages <= 0 else min(len(document), max_pages)
        image_pages: list[int] = []
        for index in range(page_count):
            page = document[index]
            try:
                if any(obj.type == raw.FPDF_PAGEOBJ_IMAGE for obj in page.get_objects()):
                    image_pages.append(index)
            finally:
                page.close()
        return image_pages
    except Exception as exc:
        warnings.append(f"PDF image detection through pypdfium2 failed: {exc}")
        return None
    finally:
        document.close()


def _ocr_pdf(
    path: Path,
    max_pages: int,
    warnings: list[str],
    page_indexes: list[int] | None = None,
    *,
    serial_ocr: bool = False,
    ocr_languages: tuple[str, ...] = (),
) -> list[str]:
    try:
        return _ocr_pdf_with_pymupdf(
            path,
            max_pages,
            warnings,
            page_indexes,
            serial_ocr=serial_ocr,
            ocr_languages=ocr_languages,
        )
    except Exception as exc:
        warnings.append(f"PDF OCR through PyMuPDF unavailable: {exc}; trying pypdfium2")
        return _ocr_pdf_with_pypdfium2(
            path,
            max_pages,
            warnings,
            page_indexes,
            serial_ocr=serial_ocr,
            ocr_languages=ocr_languages,
        )


def _ocr_image(image, page_number: int, warnings: list[str], ocr_languages: tuple[str, ...]) -> str:
    import pytesseract

    language = _ocr_language_argument(ocr_languages)
    try:
        if language is None:
            return pytesseract.image_to_string(image)
        return pytesseract.image_to_string(image, lang=language)
    except Exception as exc:
        label = language or "default"
        warnings.append(f"PDF OCR page {page_number} {label} failed: {exc}")
        return ""


def _ocr_image_serial(image, context: str, warnings: list[str], ocr_languages: tuple[str, ...]) -> list[str]:
    return [
        _ocr_image_language(image, context, language, label, warnings)
        for language, label in _serial_ocr_language_specs(ocr_languages)
    ]


def _ocr_image_language(image, context: str, language: str | None, label: str, warnings: list[str]) -> str:
    import pytesseract

    try:
        if language is None:
            return pytesseract.image_to_string(image)
        return pytesseract.image_to_string(image, lang=language)
    except Exception as exc:
        warnings.append(f"{context} OCR {label} failed: {exc}")
        return ""


def _ocr_pdf_with_pymupdf(
    path: Path,
    max_pages: int,
    warnings: list[str],
    page_indexes: list[int] | None = None,
    *,
    serial_ocr: bool = False,
    ocr_languages: tuple[str, ...] = (),
) -> list[str]:
    try:
        import fitz
        from PIL import Image
    except Exception as exc:
        raise RuntimeError(exc) from exc

    try:
        document = fitz.open(path)
    except Exception as exc:
        raise RuntimeError(f"PyMuPDF open failed: {exc}") from exc

    page_count = len(document) if max_pages <= 0 else min(len(document), max_pages)
    if page_indexes is None:
        indexes = range(page_count)
    else:
        indexes = sorted({index for index in page_indexes if 0 <= index < page_count})
    try:
        if serial_ocr:
            serial_results: list[str] = []
            for language, label in _serial_ocr_language_specs(ocr_languages):
                language_chunks: list[str] = []
                for index in indexes:
                    try:
                        page = document.load_page(index)
                        pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
                        image = Image.open(BytesIO(pixmap.tobytes("png")))
                        language_chunks.append(_ocr_image_language(image, f"PDF OCR page {index + 1}", language, label, warnings))
                    except Exception as exc:
                        warnings.append(f"PDF OCR page {index + 1} failed for {label}: {exc}")
                serial_results.append("\n".join(language_chunks))
            return serial_results

        chunks: list[str] = []
        for index in indexes:
            try:
                page = document.load_page(index)
                pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
                image = Image.open(BytesIO(pixmap.tobytes("png")))
                chunks.append(_ocr_image(image, index + 1, warnings, ocr_languages))
            except Exception as exc:
                warnings.append(f"PDF OCR page {index + 1} failed: {exc}")
        return ["\n".join(chunks)]
    finally:
        document.close()


def _ocr_pdf_with_pypdfium2(
    path: Path,
    max_pages: int,
    warnings: list[str],
    page_indexes: list[int] | None = None,
    *,
    serial_ocr: bool = False,
    ocr_languages: tuple[str, ...] = (),
) -> list[str]:
    try:
        import pypdfium2 as pdfium
    except Exception as exc:
        warnings.append(f"PDF OCR through pypdfium2 unavailable: {exc}")
        return []

    try:
        document = pdfium.PdfDocument(str(path))
    except Exception as exc:
        warnings.append(f"pypdfium2 open failed: {exc}")
        return []

    try:
        page_count = len(document) if max_pages <= 0 else min(len(document), max_pages)
        if page_indexes is None:
            indexes = range(page_count)
        else:
            indexes = sorted({index for index in page_indexes if 0 <= index < page_count})
        if serial_ocr:
            serial_results: list[str] = []
            for language, label in _serial_ocr_language_specs(ocr_languages):
                language_chunks: list[str] = []
                for index in indexes:
                    page = None
                    try:
                        page = document[index]
                        image = page.render(scale=2).to_pil()
                        language_chunks.append(_ocr_image_language(image, f"PDF OCR page {index + 1}", language, label, warnings))
                    except Exception as exc:
                        warnings.append(f"PDF OCR page {index + 1} through pypdfium2 failed for {label}: {exc}")
                    finally:
                        if page is not None:
                            page.close()
                serial_results.append("\n".join(language_chunks))
            return serial_results

        chunks: list[str] = []
        for index in indexes:
            page = None
            try:
                page = document[index]
                image = page.render(scale=2).to_pil()
                chunks.append(_ocr_image(image, index + 1, warnings, ocr_languages))
            except Exception as exc:
                warnings.append(f"PDF OCR page {index + 1} through pypdfium2 failed: {exc}")
            finally:
                if page is not None:
                    page.close()
        return ["\n".join(chunks)]
    finally:
        document.close()


def _extract_docx(path: Path, warnings: list[str]) -> str:
    try:
        from docx import Document

        document = Document(str(path))
        paragraphs = [paragraph.text for paragraph in document.paragraphs]
        table_cells = [cell.text for table in document.tables for row in table.rows for cell in row.cells]
        return "\n".join(paragraphs + table_cells)
    except Exception as exc:
        warnings.append(f"python-docx extraction failed: {exc}; trying XML fallback")
    try:
        with zipfile.ZipFile(path) as archive:
            xml = archive.read("word/document.xml").decode("utf-8", errors="ignore")
        text = re.sub(r"<[^>]+>", " ", xml)
        return html.unescape(text)
    except Exception as exc:
        warnings.append(f"DOCX XML fallback failed: {exc}")
        return ""


def _extract_table(path: Path, max_rows: int, warnings: list[str]) -> str:
    suffix = path.suffix.lower()
    if suffix == ".parquet":
        try:
            import pyarrow.parquet as pq

            table = pq.read_table(path)
            names = table.column_names
            rows = []
            for index, row in enumerate(table.to_pylist()):
                if max_rows > 0 and index >= max_rows:
                    warnings.append(f"table truncated to {max_rows} rows")
                    break
                rows.append(_format_record(row, names))
            return "\n".join(rows)
        except Exception as exc:
            warnings.append(f"Parquet extraction failed: {exc}")
            return ""
    elif suffix in {".xls", ".xlsx"}:
        return _extract_excel(path, max_rows, warnings)
    return ""


def _extract_csv(path: Path, max_rows: int, warnings: list[str]) -> str:
    text = _decode_bytes(_read_limited_bytes(path, 0, warnings))
    try:
        dialect = Sniffer().sniff(text[:4096])
    except Exception:
        dialect = "excel"
    reader = DictReader(StringIO(text), dialect=dialect)
    rows = []
    fieldnames = reader.fieldnames or []
    for index, record in enumerate(reader):
        if max_rows > 0 and index >= max_rows:
            warnings.append(f"CSV truncated to {max_rows} rows")
            break
        rows.append(_format_record(record, fieldnames))
    return "\n".join(rows)


def _format_record(record: dict, fieldnames: list[str]) -> str:
    parts = []
    for column in fieldnames:
        value = record.get(column)
        if value is None or value == "":
            continue
        parts.append(f"{column}: {value}")
    return " ".join(parts)


def _extract_excel(path: Path, max_rows: int, warnings: list[str]) -> str:
    if path.suffix.lower() == ".xlsx":
        try:
            from openpyxl import load_workbook

            workbook = load_workbook(path, read_only=True, data_only=True)
            chunks = []
            for sheet in workbook.worksheets:
                iterator = sheet.iter_rows(values_only=True)
                headers = [str(value) if value is not None else f"col_{index}" for index, value in enumerate(next(iterator, []), start=1)]
                for row_index, row in enumerate(iterator):
                    if max_rows > 0 and row_index >= max_rows:
                        warnings.append(f"sheet {sheet.title} truncated to {max_rows} rows")
                        break
                    record = {headers[index]: value for index, value in enumerate(row) if index < len(headers)}
                    chunks.append(f"sheet: {sheet.title} {_format_record(record, headers)}")
            return "\n".join(chunks)
        except Exception as exc:
            warnings.append(f"XLSX extraction failed: {exc}")
            return ""
    try:
        import xlrd

        workbook = xlrd.open_workbook(path)
        chunks = []
        for sheet in workbook.sheets():
            headers = [str(sheet.cell_value(0, col)) or f"col_{col + 1}" for col in range(sheet.ncols)]
            for row_index in range(1, sheet.nrows):
                if max_rows > 0 and row_index > max_rows:
                    warnings.append(f"sheet {sheet.name} truncated to {max_rows} rows")
                    break
                record = {headers[col]: sheet.cell_value(row_index, col) for col in range(sheet.ncols)}
                chunks.append(f"sheet: {sheet.name} {_format_record(record, headers)}")
        return "\n".join(chunks)
    except Exception as exc:
        warnings.append(f"XLS extraction failed: {exc}")
        return ""


def _extract_image(path: Path, warnings: list[str], ocr_languages: tuple[str, ...]) -> str:
    try:
        from PIL import Image
        import pytesseract

        image = Image.open(path)
        language = _ocr_language_argument(ocr_languages)
        try:
            if language is None:
                return pytesseract.image_to_string(image)
            return pytesseract.image_to_string(image, lang=language)
        except Exception as exc:
            label = language or "default"
            warnings.append(f"OCR {label} failed: {exc}")
            return ""
    except Exception as exc:
        warnings.append(f"OCR failed: {exc}")
        return ""


def _extract_image_serial(path: Path, warnings: list[str], ocr_languages: tuple[str, ...]) -> list[str]:
    try:
        from PIL import Image

        image = Image.open(path)
        specs = ", ".join(label for _, label in _serial_ocr_language_specs(ocr_languages))
        warnings.append(f"serial OCR enabled for image: {specs}")
        return _ocr_image_serial(image, f"OCR image {path.name}", warnings, ocr_languages)
    except Exception as exc:
        warnings.append(f"OCR failed: {exc}")
        return []


def _extract_binary_strings(path: Path, max_bytes: int, warnings: list[str]) -> str:
    data = _read_limited_bytes(path, max_bytes, warnings)
    if not _looks_like_text_bytes(data):
        warnings.append("binary file skipped; unsupported binary format")
        return ""
    decoded = _decode_bytes(data)
    chunks = re.findall(r"[A-Za-zА-Яа-яЁё0-9@._<>\-+/\\,;:() ]{5,}", decoded)
    return "\n".join(chunks)


def extract_text(
    path: Path,
    *,
    enable_ocr: bool = False,
    serial_ocr: bool = False,
    ocr_languages: tuple[str, ...] | list[str] | None = None,
    max_bytes: int = 20_000_000,
    max_chars: int = 2_000_000,
    max_rows: int = 200_000,
    max_pdf_pages: int = 0,
) -> ExtractionResult:
    warnings: list[str] = []
    normalized_ocr_languages = normalize_ocr_languages(ocr_languages)
    suffix = path.suffix.lower()
    text = ""
    scan_texts: list[str] = []
    try:
        if suffix == ".csv":
            text = _extract_csv(path, max_rows, warnings)
        elif suffix in TEXT_EXTENSIONS:
            text = _decode_bytes(_read_limited_bytes(path, max_bytes, warnings))
        elif suffix in HTML_EXTENSIONS:
            raw = _decode_bytes(_read_limited_bytes(path, max_bytes, warnings))
            text = _strip_html(raw)
        elif suffix == ".rtf":
            raw = _decode_bytes(_read_limited_bytes(path, max_bytes, warnings))
            text = _strip_rtf(raw)
        elif suffix == ".pdf":
            text, scan_texts = _extract_pdf(
                path,
                max_pdf_pages,
                enable_ocr,
                serial_ocr,
                normalized_ocr_languages,
                warnings,
            )
        elif suffix == ".docx":
            text = _extract_docx(path, warnings)
        elif suffix in {".parquet", ".xls", ".xlsx"}:
            text = _extract_table(path, max_rows, warnings)
        elif suffix in IMAGE_EXTENSIONS:
            if enable_ocr:
                if serial_ocr:
                    scan_texts = [part for part in _extract_image_serial(path, warnings, normalized_ocr_languages) if part.strip()]
                    text = "\n".join(scan_texts)
                else:
                    text = _extract_image(path, warnings, normalized_ocr_languages)
            else:
                warnings.append("image skipped; pass --ocr or --serial-ocr to enable OCR")
        elif suffix in {".doc"}:
            warnings.append("legacy binary office file processed with string fallback")
            text = _extract_binary_strings(path, max_bytes, warnings)
        else:
            warnings.append("unsupported extension processed as text only if bytes look textual")
            text = _extract_binary_strings(path, max_bytes, warnings)
    except Exception as exc:
        warnings.append(f"extraction failed: {exc}")
        text = ""
        scan_texts = []
    return ExtractionResult(
        text=_limit_text(text, max_chars, warnings),
        file_format=suffix or "[no_ext]",
        warnings=warnings,
        scan_texts=[_limit_text(part, max_chars, warnings) for part in scan_texts],
    )
